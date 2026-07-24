from __future__ import annotations

import math
from datetime import date, datetime
from pathlib import Path
from typing import Optional

import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from backend.services.llm import generate_cover_letter_body, generate_cv_summary
from core.config import LLMConfig, StorageConfig, UserProfile
from core.country_styles import CountryStyle, get_style
from core.storage import get_db_connection

# ── DOCX helpers ──────────────────────────────────────────────────────────────


def _set_doc_defaults(doc: Document, style: CountryStyle) -> None:
    for section in doc.sections:
        m = Cm(style.margin_cm)
        section.top_margin = m
        section.bottom_margin = m
        section.left_margin = m
        section.right_margin = m
    normal = doc.styles["Normal"]
    normal.font.name = style.font_family
    normal.font.size = Pt(style.font_size_pt)


def _add_bottom_border(paragraph) -> None:
    """Add a thin bottom border to a paragraph (used for section dividers)."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _add_name_header(doc: Document, name: str, contact_line: str, style: CountryStyle) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(name)
    run.bold = True
    run.font.name = style.font_family
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(8)
    run2 = p2.add_run(contact_line)
    run2.font.name = style.font_family
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def _add_section_header(doc: Document, title: str, style: CountryStyle) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title.upper())
    run.bold = True
    run.font.name = style.font_family
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    _add_bottom_border(p)


def _add_body_para(doc: Document, text: str, style: CountryStyle) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = style.font_family
    run.font.size = Pt(style.font_size_pt)


def _add_bullet(doc: Document, text: str, style: CountryStyle) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = style.font_family
    run.font.size = Pt(style.font_size_pt)


def _add_experience_entry(doc: Document, exp, style: CountryStyle) -> None:
    """Add one work-experience block: role/company on one line, bullets below."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    title_run = p.add_run(f"{exp.title}  •  {exp.company}")
    title_run.bold = True
    title_run.font.name = style.font_family
    title_run.font.size = Pt(style.font_size_pt)

    location_part = f"  |  {exp.location}" if exp.location else ""
    date_run = p.add_run(f"\n{exp.start_date} – {exp.end_date}{location_part}")
    date_run.font.name = style.font_family
    date_run.font.size = Pt(9)
    date_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    for bullet in exp.bullets:
        _add_bullet(doc, bullet, style)


# ── Cover letter assembler ────────────────────────────────────────────────────


def build_cover_letter(
    path: Path,
    profile: UserProfile,
    job_title: str,
    company: str,
    body_text: str,
    style: CountryStyle,
) -> None:
    doc = Document()
    _set_doc_defaults(doc, style)

    # Contact info
    contact_parts = []
    if profile.personal.location:
        contact_parts.append(profile.personal.location)
    contact_parts.append("girireddy@multicase.com")  # placeholder; user can update profile
    _add_name_header(doc, profile.personal.name, "  |  ".join(contact_parts), style)

    # Date
    if style.doc_language == "de":
        date_str = date.today().strftime("%-d. %B %Y")
    else:
        date_str = date.today().strftime("%B %-d, %Y")
    p_date = doc.add_paragraph()
    p_date.paragraph_format.space_after = Pt(10)
    p_date.add_run(date_str).font.size = Pt(style.font_size_pt)

    # Recipient block (Germany-style)
    if style.include_address_block:
        p_recip = doc.add_paragraph()
        p_recip.paragraph_format.space_after = Pt(10)
        run_recip = p_recip.add_run(f"{company}\nHR Department")
        run_recip.font.size = Pt(style.font_size_pt)

    # Subject line
    p_subj = doc.add_paragraph()
    p_subj.paragraph_format.space_after = Pt(10)
    subj_run = p_subj.add_run(f"{style.cl_subject_prefix} {job_title}")
    subj_run.bold = True
    subj_run.font.size = Pt(style.font_size_pt)

    # Greeting
    _add_body_para(doc, style.cl_greeting, style)

    # Body paragraphs
    for para in body_text.split("\n\n"):
        para = para.strip()
        if para:
            _add_body_para(doc, para, style)

    # Closing
    doc.add_paragraph()
    p_close = doc.add_paragraph()
    p_close.paragraph_format.space_after = Pt(24)
    p_close.add_run(style.cl_closing).font.size = Pt(style.font_size_pt)

    p_sig = doc.add_paragraph()
    sig_run = p_sig.add_run(profile.personal.name)
    sig_run.bold = True
    sig_run.font.size = Pt(style.font_size_pt)

    doc.save(str(path))


# ── CV assembler ──────────────────────────────────────────────────────────────


def build_cv(
    path: Path,
    profile: UserProfile,
    summary_text: str,
    style: CountryStyle,
) -> None:
    headers = style.cv_section_headers
    doc = Document()
    _set_doc_defaults(doc, style)

    # Header
    contact_parts = []
    if profile.personal.location:
        contact_parts.append(profile.personal.location)
    contact_parts.append("girireddy@multicase.com")
    _add_name_header(doc, profile.personal.name, "  |  ".join(contact_parts), style)

    # Summary / Profile
    _add_section_header(doc, headers.get("summary", "Profile"), style)
    _add_body_para(doc, summary_text, style)

    # Experience
    if profile.experience:
        _add_section_header(doc, headers.get("experience", "Work Experience"), style)
        for exp in profile.experience:
            _add_experience_entry(doc, exp, style)

    # Education
    if profile.education:
        _add_section_header(doc, headers.get("education", "Education"), style)
        for edu in profile.education:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            line1 = f"{edu.degree} in {edu.field}" if edu.field else edu.degree
            edu_run = p.add_run(line1)
            edu_run.bold = True
            edu_run.font.size = Pt(style.font_size_pt)
            line2 = edu.institution if edu.institution else "[Institution]"
            if edu.country:
                line2 += f", {edu.country}"
            if edu.year:
                line2 += f"  ({edu.year})"
            date_run = p.add_run(f"\n{line2}")
            date_run.font.size = Pt(9)
            date_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Skills
    all_skill_groups = [
        ("Programming", profile.skills.programming),
        ("ML / AI", profile.skills.ml_frameworks + profile.skills.ai_llm),
        ("Cloud", profile.skills.cloud),
        ("Data Engineering", profile.skills.data_engineering),
        ("Tools", profile.skills.tools),
    ]
    skill_groups = [(label, skills) for label, skills in all_skill_groups if skills]
    if skill_groups:
        _add_section_header(doc, headers.get("skills", "Skills"), style)
        for label, skills in skill_groups:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(1)
            label_run = p.add_run(f"{label}: ")
            label_run.bold = True
            label_run.font.size = Pt(style.font_size_pt)
            skills_run = p.add_run(", ".join(skills))
            skills_run.font.size = Pt(style.font_size_pt)

    # Languages
    if profile.languages:
        _add_section_header(doc, headers.get("languages", "Languages"), style)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        lang_strs = [f"{e.language} ({e.level})" for e in profile.languages]
        p.add_run(", ".join(lang_strs)).font.size = Pt(style.font_size_pt)

    doc.save(str(path))


# ── DocumentService ───────────────────────────────────────────────────────────


class DocumentService:
    def __init__(
        self,
        storage: StorageConfig,
        llm: LLMConfig,
        profile: UserProfile,
    ) -> None:
        self._storage = storage
        self._llm = llm
        self._profile = profile

    def generate_and_save(
        self,
        job_uuid: str,
        doc_type: str,
        job_row: pd.Series,
    ) -> None:
        """Generate DOCX files and persist their paths to SQLite.

        Designed to run as a FastAPI BackgroundTask — errors are logged, not raised.
        """
        try:
            country = str(job_row.get("country") or "")
            job_title = str(job_row.get("title") or "this position")
            company = str(job_row.get("company") or "your company")
            description = str(job_row.get("description") or "")[:1000]
            style = get_style(country)

            out_dir = Path(self._storage.documents_dir) / job_uuid
            out_dir.mkdir(parents=True, exist_ok=True)

            now = datetime.utcnow().isoformat()
            records: list[tuple] = []

            if doc_type in ("cover_letter", "both"):
                body = generate_cover_letter_body(
                    self._llm, self._profile, job_title, company, description, style
                )
                cl_path = out_dir / "cover_letter.docx"
                build_cover_letter(cl_path, self._profile, job_title, company, body, style)
                records.append(("cover_letter", str(cl_path), country, now))

            if doc_type in ("cv", "both"):
                summary = generate_cv_summary(
                    self._llm, self._profile, job_title, company, description, style
                )
                cv_path = out_dir / "cv.docx"
                build_cv(cv_path, self._profile, summary, style)
                records.append(("cv", str(cv_path), country, now))

            with get_db_connection(self._storage) as conn:
                # Ensure job is tracked in job_status (required by FK)
                conn.execute(
                    """
                    INSERT OR IGNORE INTO job_status (uuid, status, updated_at)
                    VALUES (?, 'interesting', ?)
                    """,
                    (job_uuid, now),
                )
                for doc_type_val, file_path, country_style, created_at in records:
                    conn.execute(
                        """
                        INSERT INTO documents
                            (job_uuid, doc_type, file_path, country_style, created_at)
                        VALUES (?, ?, ?, ?, ?)
                        """,
                        (job_uuid, doc_type_val, file_path, country_style, created_at),
                    )
                conn.commit()

            print(
                f"[documents] Generated {len(records)} file(s) for job {job_uuid} "
                f"({country}, {doc_type})"
            )
        except Exception as exc:
            print(f"[documents] Generation failed for {job_uuid}: {exc}")
