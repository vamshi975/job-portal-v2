"""
Tests for document generation:
- DOCX files are well-formed (openable, have expected content)
- LLM prompt helpers produce non-empty output with correct language
- Country style lookup covers all 8 configured countries
- DocumentService.generate_and_save writes files and SQLite rows
"""
from __future__ import annotations

import tempfile
from pathlib import Path
from unittest.mock import MagicMock, patch

import pandas as pd
import pytest

from core.config import (
    AppConfig,
    BackendConfig,
    ExperienceEntry,
    LanguageEntry,
    LLMConfig,
    PersonalInfo,
    SkillsConfig,
    StorageConfig,
    UserProfile,
)
from core.country_styles import get_style
from core.storage import ensure_db_schema, get_db_connection


# ── fixtures ──────────────────────────────────────────────────────────────────


@pytest.fixture
def profile() -> UserProfile:
    return UserProfile(
        personal=PersonalInfo(
            name="Vamshi Bhushanaboina",
            current_title="Data Scientist",
            experience_years=5,
            location="Germany",
        ),
        skills=SkillsConfig(
            programming=["Python", "SQL"],
            ml_frameworks=["scikit-learn", "PyTorch"],
            cloud=["Azure", "Azure ML"],
            ai_llm=["RAG", "LLM", "LangChain"],
            data_engineering=["PySpark", "dbt"],
            tools=["Docker", "Git"],
        ),
        experience=[
            ExperienceEntry(
                title="Data Scientist",
                company="Acme GmbH",
                location="Berlin",
                start_date="2022-01",
                end_date="present",
                bullets=["Built RAG pipeline", "Deployed Azure ML model"],
            )
        ],
        languages=[
            LanguageEntry(language="English", level="C1"),
            LanguageEntry(language="German", level="B2"),
        ],
        target_roles=["Data Scientist", "ML Engineer"],
        industries=["Insurance", "CRM"],
    )


@pytest.fixture
def tmp_storage(tmp_path: Path) -> StorageConfig:
    return StorageConfig(
        daily_data_dir=str(tmp_path / "daily_data"),
        processed_data_dir=str(tmp_path / "data"),
        documents_dir=str(tmp_path / "documents"),
        db_path=str(tmp_path / "data" / "jobs.db"),
    )


@pytest.fixture
def sample_job_row() -> pd.Series:
    return pd.Series({
        "uuid": "test-uuid-123",
        "title": "Data Scientist",
        "company": "Acme GmbH",
        "country": "Germany",
        "description": "We need Python, ML, and Azure experience.",
    })


# ── country style tests ────────────────────────────────────────────────────────


def test_all_eight_countries_have_styles():
    countries = [
        "Germany", "Netherlands", "Belgium", "India",
        "United States", "Canada", "Denmark", "Sweden",
    ]
    for country in countries:
        style = get_style(country)
        assert style.font_family, f"{country} has no font_family"
        assert style.cl_greeting, f"{country} has no cl_greeting"
        assert style.cv_section_headers, f"{country} has no cv_section_headers"


def test_germany_is_german_language():
    style = get_style("Germany")
    assert style.doc_language == "de"


def test_unknown_country_returns_fallback():
    style = get_style("Mars")
    assert style.doc_language == "en"
    assert style.font_family  # not empty


def test_germany_has_address_block():
    style = get_style("Germany")
    assert style.include_address_block is True


def test_us_has_no_address_block():
    style = get_style("United States")
    assert style.include_address_block is False


# ── DOCX builder tests ────────────────────────────────────────────────────────


def test_build_cover_letter_creates_file(tmp_path, profile):
    from backend.services.documents import build_cover_letter

    style = get_style("Germany")
    out = tmp_path / "cl.docx"
    build_cover_letter(out, profile, "Data Scientist", "Acme GmbH", "Body text here.", style)
    assert out.exists()
    assert out.stat().st_size > 0


def test_build_cv_creates_file(tmp_path, profile):
    from backend.services.documents import build_cv

    style = get_style("Germany")
    out = tmp_path / "cv.docx"
    build_cv(out, profile, "Experienced data scientist specialising in ML.", style)
    assert out.exists()
    assert out.stat().st_size > 0


def test_cover_letter_contains_name(tmp_path, profile):
    from backend.services.documents import build_cover_letter
    from docx import Document

    style = get_style("Netherlands")
    out = tmp_path / "cl.docx"
    build_cover_letter(out, profile, "ML Engineer", "DataCo", "Body text.", style)

    doc = Document(str(out))
    full_text = " ".join(p.text for p in doc.paragraphs)
    assert "Vamshi" in full_text


def test_cv_contains_experience_bullet(tmp_path, profile):
    from backend.services.documents import build_cv
    from docx import Document

    style = get_style("Germany")
    out = tmp_path / "cv.docx"
    build_cv(out, profile, "Erfahrener Data Scientist.", style)

    doc = Document(str(out))
    full_text = " ".join(p.text for p in doc.paragraphs)
    assert "Built RAG pipeline" in full_text


# ── LLM placeholder tests ──────────────────────────────────────────────────────


def test_cl_placeholder_returned_when_llm_not_configured(profile):
    from backend.services.llm import generate_cover_letter_body

    llm = LLMConfig(base_url="", model="")
    style = get_style("Germany")
    result = generate_cover_letter_body(llm, profile, "Data Scientist", "Acme", "Desc", style)
    assert len(result) > 50
    assert "Vamshi" not in result  # placeholder doesn't use name in body


def test_cv_summary_placeholder_returned_when_not_configured(profile):
    from backend.services.llm import generate_cv_summary

    llm = LLMConfig(base_url="", model="")
    style = get_style("United States")
    result = generate_cv_summary(llm, profile, "Data Scientist", "Acme", "Desc", style)
    assert len(result) > 30


def test_cl_uses_llm_when_configured(profile):
    from backend.services.llm import generate_cover_letter_body

    llm = LLMConfig(base_url="https://example.com/v1/", model="gemma-4-31b-it")
    style = get_style("Netherlands")

    mock_choice = MagicMock()
    mock_choice.message.content = "This is a generated cover letter body.\n\nSecond paragraph."
    mock_resp = MagicMock()
    mock_resp.choices = [mock_choice]

    with patch("openai.OpenAI") as mock_cls:
        mock_client = MagicMock()
        mock_client.chat.completions.create.return_value = mock_resp
        mock_cls.return_value = mock_client

        result = generate_cover_letter_body(llm, profile, "ML Engineer", "DataCo", "Desc", style)

    assert "generated cover letter" in result


def test_cl_falls_back_on_llm_error(profile):
    from backend.services.llm import generate_cover_letter_body

    llm = LLMConfig(base_url="https://example.com/v1/", model="gemma-4-31b-it")
    style = get_style("Germany")

    with patch("openai.OpenAI") as mock_cls:
        mock_client = MagicMock()
        mock_client.chat.completions.create.side_effect = RuntimeError("network error")
        mock_cls.return_value = mock_client

        result = generate_cover_letter_body(llm, profile, "DS", "Co", "Desc", style)

    assert len(result) > 20  # returns placeholder


# ── DocumentService integration tests ────────────────────────────────────────


def test_generate_and_save_writes_files_and_db(tmp_path, profile, sample_job_row):
    from backend.services.documents import DocumentService

    storage = StorageConfig(
        daily_data_dir=str(tmp_path / "daily_data"),
        processed_data_dir=str(tmp_path / "data"),
        documents_dir=str(tmp_path / "documents"),
        db_path=str(tmp_path / "data" / "jobs.db"),
    )
    ensure_db_schema(storage)

    llm = LLMConfig(base_url="", model="")  # uses placeholders
    svc = DocumentService(storage, llm, profile)
    svc.generate_and_save("test-uuid-123", "both", sample_job_row)

    # Files on disk
    doc_dir = Path(storage.documents_dir) / "test-uuid-123"
    assert (doc_dir / "cover_letter.docx").exists()
    assert (doc_dir / "cv.docx").exists()

    # SQLite rows
    with get_db_connection(storage) as conn:
        rows = conn.execute(
            "SELECT doc_type FROM documents WHERE job_uuid = ?", ("test-uuid-123",)
        ).fetchall()
    doc_types = {row["doc_type"] for row in rows}
    assert doc_types == {"cover_letter", "cv"}


def test_generate_cover_letter_only(tmp_path, profile, sample_job_row):
    from backend.services.documents import DocumentService

    storage = StorageConfig(
        daily_data_dir=str(tmp_path / "daily_data"),
        processed_data_dir=str(tmp_path / "data"),
        documents_dir=str(tmp_path / "documents"),
        db_path=str(tmp_path / "data" / "jobs.db"),
    )
    ensure_db_schema(storage)

    svc = DocumentService(storage, LLMConfig(base_url="", model=""), profile)
    svc.generate_and_save("test-uuid-cl", "cover_letter", sample_job_row)

    doc_dir = Path(storage.documents_dir) / "test-uuid-cl"
    assert (doc_dir / "cover_letter.docx").exists()
    assert not (doc_dir / "cv.docx").exists()

    with get_db_connection(storage) as conn:
        count = conn.execute(
            "SELECT COUNT(*) FROM documents WHERE job_uuid = ?", ("test-uuid-cl",)
        ).fetchone()[0]
    assert count == 1
